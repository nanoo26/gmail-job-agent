"""
cv_loader.py – טוען קורות חיים מ-DOCX ו-PDF ומחזיר טוקנים לחישוב cv_boost.

נטען מחדש בכל הפעלה → שינוי קובץ CV ישתקף אוטומטית ב-cv_boost.
כל מסלול יכול לכלול מספר קבצים (DOCX + PDF); הטוקנים מאוחדים (union).
"""

import os
import re
import sys
from pathlib import Path

if hasattr(sys.stdout, "reconfigure"):
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")

# ── תיקיית קורות חיים ────────────────────────────────────────────────
# ניתן להגדיר CV_DIR כמשתנה סביבה; ברירת המחדל היא תת-תיקייה "cvs"
# בספריית הפרויקט.  דוגמה ב-.env:  CV_DIR=C:\Users\me\Documents\cvs
_CV_DIR = os.getenv("CV_DIR", str(Path(__file__).parent / "cvs"))

# ── מיפוי מסלולים → קבצים (DOCX ו/או PDF) ───────────────────────────
# כל ערך הוא רשימה של נתיבים; ניתן לשלב DOCX + PDF לאותו מסלול.
# הטוקנים מכל הקבצים מאוחדים ← token-pool רחב יותר = boost יותר מדויק.
CV_FILES: dict[str, list[str]] = {
    "IT": [
        str(Path(_CV_DIR) / "קורות חיים - IT.docx"),
        str(Path(_CV_DIR) / "קורות חיים - IT.pdf"),
    ],
    "תפעול": [
        str(Path(_CV_DIR) / "קורות חיים - תפעול.docx"),
        str(Path(_CV_DIR) / "קורות חיים - תפעול.pdf"),
    ],
    "אחזקה": [
        str(Path(_CV_DIR) / "קורות חיים - אחזקה.docx"),
    ],
}

# ── מילות עצירה ──────────────────────────────────────────────────────
STOPWORDS = {
    # אנגלית
    "the", "and", "or", "in", "on", "at", "to", "of", "for", "with",
    "by", "from", "is", "are", "was", "were", "be", "been", "that",
    "this", "it", "as", "not", "but", "have", "has", "had", "will",
    "would", "can", "could", "should", "may", "might", "do", "did",
    "does", "more", "my", "we", "you", "he", "she", "his", "her",
    "our", "their", "its", "all", "some", "any", "one", "two", "three",
    # עברית
    "של", "את", "עם", "על", "אל", "לא", "כי", "הם", "הן", "הוא", "היא",
    "אני", "אנחנו", "אתם", "אתן", "זה", "זו", "כל", "יש", "אין", "עד",
    "כן", "גם", "רק", "אחרי", "לפני", "בין", "מתוך", "תוך", "היה",
    "הייתה", "הייתי", "יהיה", "תהיה", "כבר", "עוד", "אבל", "אם", "כך",
}

MIN_TOKEN_LEN = 3
CV_BOOST_CAP  = 20


# ── tokenize ─────────────────────────────────────────────────────────
def tokenize(text: str) -> set:
    """מפרק טקסט למילים ייחודיות (lowercase, אנגלית + עברית, ללא stopwords)."""
    words = re.findall(r"[a-z\u05d0-\u05ea0-9]{%d,}" % MIN_TOKEN_LEN, text.lower())
    return {w for w in words if w not in STOPWORDS}


# ── DOCX reader ───────────────────────────────────────────────────────
def load_cv_docx(path: str) -> str:
    """קורא DOCX ומחזיר טקסט (פסקאות + תאי טבלה)."""
    try:
        from docx import Document
        doc = Document(path)
        parts = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    parts.append(cell.text)
        return "\n".join(parts)
    except FileNotFoundError:
        return ""
    except Exception as e:
        print(f"  [cv_loader] שגיאה (DOCX) {Path(path).name}: {e}")
        return ""


# ── PDF reader ────────────────────────────────────────────────────────
def load_cv_pdf(path: str) -> str:
    """
    קורא PDF ומחזיר טקסט מלא.
    משתמש ב-pdfplumber (מדויק יותר לקו"ח מעוצבים, כולל RTL).
    Fallback ל-pypdf אם pdfplumber לא מותקן.
    """
    try:
        import pdfplumber
        with pdfplumber.open(path) as pdf:
            pages = [page.extract_text() or "" for page in pdf.pages]
        return "\n".join(pages)
    except ImportError:
        pass
    except FileNotFoundError:
        return ""
    except Exception as e:
        print(f"  [cv_loader] שגיאה (PDF/pdfplumber) {Path(path).name}: {e}")
        return ""

    # fallback: pypdf
    try:
        from pypdf import PdfReader
        reader = PdfReader(path)
        return "\n".join(page.extract_text() or "" for page in reader.pages)
    except FileNotFoundError:
        return ""
    except Exception as e:
        print(f"  [cv_loader] שגיאה (PDF/pypdf) {Path(path).name}: {e}")
        return ""


def _load_any(path: str) -> str:
    """בורר DOCX / PDF לפי סיומת."""
    ext = Path(path).suffix.lower()
    if ext == ".docx":
        return load_cv_docx(path)
    if ext == ".pdf":
        return load_cv_pdf(path)
    return ""


# ── load_all_cvs ──────────────────────────────────────────────────────
def load_all_cvs(base_dir: str = ".") -> dict:
    """
    טוען את כל קורות החיים ומחזיר: { track: set_of_tokens }

    נטען מחדש בכל הפעלה — שינוי קובץ CV ישתקף אוטומטית.
    טוקנים מכל קבצי המסלול מאוחדים (union).
    """
    print("\n" + "=" * 62)
    print("  טעינת קורות חיים (נלמד מחדש בכל הפעלה)")
    print("=" * 62)

    result: dict[str, set] = {}
    ok_tracks = 0

    for track, file_list in CV_FILES.items():
        combined_tokens: set = set()
        file_results = []

        for filepath in file_list:
            full_path = (
                filepath if Path(filepath).is_absolute()
                else str(Path(base_dir) / filepath)
            )
            text = _load_any(full_path)
            fname = Path(full_path).name
            if text.strip():
                tokens = tokenize(text)
                combined_tokens |= tokens          # ← union
                file_results.append(f"    ✔ {fname}  ({len(tokens)} טוקנים)")
            else:
                file_results.append(f"    ✘ {fname}  (לא נמצא / ריק)")

        result[track] = combined_tokens
        status = "✔" if combined_tokens else "✘"
        print(f"\n  [{status}] {track}  —  {len(combined_tokens)} טוקנים ייחודיים (union)")
        for line in file_results:
            print(line)

        if combined_tokens:
            ok_tracks += 1

    print("\n" + "-" * 62)
    print(f"  סה\"כ: {ok_tracks}/{len(CV_FILES)} מסלולים עם טוקנים")
    print("=" * 62 + "\n")
    return result


# ── compute_cv_boost ──────────────────────────────────────────────────
def compute_cv_boost(cv_tokens: set, email_text: str, top_n: int = 10) -> tuple:
    """
    מחשב חפיפה בין טוקני קורות החיים (union כל הקבצים) לטקסט המייל.
    מחזיר (cv_boost: int 0–CV_BOOST_CAP, match_reasons: str)
    """
    if not cv_tokens:
        return 0, ""

    email_tokens = tokenize(email_text)
    overlap = cv_tokens & email_tokens
    sorted_overlap = sorted(overlap, key=len, reverse=True)[:top_n]

    boost = min(CV_BOOST_CAP, len(overlap) * 2)
    reasons = ", ".join(sorted_overlap) if sorted_overlap else ""
    return boost, reasons
